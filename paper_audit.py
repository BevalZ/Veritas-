#!/usr/bin/env python3
"""Paper Audit - 学术论文自动审查工具 [耿同学版]
基于3个开源项目思路开发：
- wooly99/geng-academic-fraud-detector 耿同学六式
- NeoSpecies/AcademicIntegrityHunter 本地统计算法
- jingshouyan/academic-integrity-geng 五维审查体系
输入PDF路径 → MinerU转Markdown → 本地统计检测 + LLM语义分析 → 输出md格式报告
用法: python paper_audit.py <pdf_path> [--mineru] [--max-chars 8000] [--output report.md]
"""
import re, json, time, argparse, urllib.request, zlib, math, collections, os, mimetypes, fnmatch, csv, platform, webbrowser, subprocess, sys, requests, builtins, hashlib
from pathlib import Path
from typing import Tuple, Dict, List

# Windows/重定向控制台默认GBK时，emoji/中文符号可能触发UnicodeEncodeError；统一兜底为UTF-8。
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


# ══════════════════════════════════════════════════════════════
# 运行日志 / 进度条 / 输出目录控制
# ══════════════════════════════════════════════════════════════
_ORIGINAL_PRINT = builtins.print
_RUN_LOG_FILE = None
_RUN_OUTPUT_DIR = None
_RUN_OUTPUT_STEM = None


def get_output_base(input_path: Path):
    """返回所有运行产物的基准目录和名称。

    规则：输入文件→文件所在目录；输入目录→该目录本身。
    """
    input_path = Path(input_path)
    if input_path.is_dir():
        return input_path, input_path.name or "audit_report"
    return input_path.parent, input_path.stem


def _safe_name(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*\x00-\x1f]+', '_', str(name)).strip(' .') or 'paper_audit'


def setup_run_logging(input_path: Path):
    """把print同时写到控制台和同目录log文件。"""
    global _RUN_LOG_FILE, _RUN_OUTPUT_DIR, _RUN_OUTPUT_STEM
    out_dir, stem = get_output_base(Path(input_path))
    out_dir.mkdir(parents=True, exist_ok=True)
    _RUN_OUTPUT_DIR = out_dir
    _RUN_OUTPUT_STEM = _safe_name(stem)
    _RUN_LOG_FILE = out_dir / f"{_RUN_OUTPUT_STEM}.paper_audit.log"
    _RUN_LOG_FILE.write_text(
        f"Paper Audit Log\nSTART {time.strftime('%F %T')}\nINPUT {Path(input_path)}\nOUTPUT_DIR {out_dir}\n\n",
        encoding="utf-8"
    )

    def tee_print(*args, **kwargs):
        _ORIGINAL_PRINT(*args, **kwargs)
        try:
            sep = kwargs.get("sep", " ")
            end = kwargs.get("end", "\n")
            msg = sep.join(str(a) for a in args) + end
            with _RUN_LOG_FILE.open("a", encoding="utf-8", errors="replace") as f:
                f.write(msg)
        except Exception:
            pass
    builtins.print = tee_print
    print(f"🧾 日志文件: {_RUN_LOG_FILE}")
    return _RUN_LOG_FILE


def get_resume_dir(output_dir: Path, output_stem: str):
    """断点续作缓存目录。"""
    d = Path(output_dir) / f".{_safe_name(output_stem)}.paper_audit_resume"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _json_load(path: Path, default=None):
    try:
        if Path(path).exists():
            return json.loads(Path(path).read_text(encoding="utf-8"))
    except Exception as e:
        print(f"⚠️ 读取缓存失败 {path}: {e}")
    return default


def _json_save(path: Path, data):
    try:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_suffix(path.suffix + ".tmp")
        tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)
    except Exception as e:
        print(f"⚠️ 写入缓存失败 {path}: {e}")


def resume_event(resume_dir: Path, step: str, status: str, detail: str = "", **extra):
    """记录可断点续作的步骤清单，同时写入普通log。"""
    try:
        event = {"time": time.strftime("%F %T"), "step": step, "status": status, "detail": detail}
        event.update(extra)
        manifest = Path(resume_dir) / "resume_manifest.jsonl"
        manifest.parent.mkdir(parents=True, exist_ok=True)
        with manifest.open("a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
        print(f"🧭 断点记录: {step} | {status} | {detail}")
    except Exception as e:
        print(f"⚠️ 写入断点记录失败: {e}")


def _text_fingerprint(text: str, extra: str = ""):
    h = hashlib.sha256()
    h.update((text or "").encode("utf-8", errors="ignore"))
    h.update(str(extra).encode("utf-8", errors="ignore"))
    return h.hexdigest()[:16]


# LLM运行参数：由CLI覆盖。默认保守，避免一次请求无限阻塞。
LLM_TIMEOUT = 45
LLM_RETRIES = 1


def load_mykey_llm_config(provider_name: str = None, ignore_config: bool = False):
    """从 ../mykey.py 或 ./mykey.py 覆盖LLM配置；不打印密钥。provider_name如 native_oai_config_3。"""
    global LLM_API_KEY, LLM_API_URL, LLM_MODEL
    try:
        _sys_path_backup = sys.path[:]
        sys.path.insert(0, str(Path(__file__).parent.parent))
        sys.path.insert(0, str(Path(__file__).parent))
        mykey = importlib.import_module("mykey")
        candidates = [provider_name] if provider_name else ["native_oai_config_0", "native_oai_config_3", "native_oai_config_2", "native_oai_config_1"]
        for cfg_name in candidates:
            cfg = getattr(mykey, cfg_name, None)
            if isinstance(cfg, dict) and cfg.get("apikey") and cfg.get("apibase"):
                LLM_API_KEY = cfg["apikey"]
                apibase = cfg.get("apibase", "")
                LLM_API_URL = f"{apibase}/chat/completions" if "/chat/completions" not in apibase else apibase
                LLM_MODEL = cfg.get("model", LLM_MODEL)
                print(f"✅ 已覆盖LLM供应商: mykey.py.{cfg_name} (模型: {LLM_MODEL}, URL: {LLM_API_URL})")
                return True
        print(f"⚠️ mykey.py中未找到可用LLM配置: {candidates}")
    except Exception as e:
        print(f"⚠️ 覆盖LLM供应商失败: {e}")
    finally:
        if '_sys_path_backup' in locals():
            sys.path = _sys_path_backup
    return False


def progress_bar(current, total, label="", width=28):
    """打印一行文本进度条；日志中保留每次更新。"""
    try:
        total = max(int(total), 1)
        current = max(0, min(int(current), total))
        filled = int(width * current / total)
        bar = "█" * filled + "░" * (width - filled)
        pct = current * 100 / total
        print(f"📊 [{bar}] {current}/{total} {pct:5.1f}% {label}")
    except Exception:
        print(f"📊 {current}/{total} {label}")


def save_mineru_artifacts(zip_url: str, zip_data: bytes, source_name: str, output_dir=None, batch_id=None):
    """把MinerU下载链接和zip保存到与输入文件/目录一致的位置。"""
    out_dir = Path(output_dir) if output_dir else (_RUN_OUTPUT_DIR or Path.cwd())
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_name(Path(source_name).stem if source_name else (batch_id or "mineru"))
    suffix = f".{_safe_name(batch_id)}" if batch_id else ""
    link_path = out_dir / f"{stem}{suffix}.mineru_url.txt"
    zip_path = out_dir / f"{stem}{suffix}.mineru.zip"
    link_path.write_text(zip_url + "\n", encoding="utf-8")
    zip_path.write_bytes(zip_data)
    print(f"  🔗 MinerU下载链接已保存: {link_path}")
    print(f"  📦 MinerU zip已保存: {zip_path} ({len(zip_data)/1024/1024:.2f}MB)")
    return zip_path, link_path

# 可选依赖：处理Word/Excel/Supplement文件
try:
    from docx import Document
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False
try:
    from openpyxl import load_workbook
    EXCEL_SUPPORTED = True
except ImportError:
    EXCEL_SUPPORTED = False

# ══════════════════════════════════════════════════════════════
# 配置区
# ══════════════════════════════════════════════════════════════
import importlib

# 尝试加载外部配置文件
LLM_API_KEY = ""
LLM_API_URL = "https://api.openai.com/v1/chat/completions"
LLM_MODEL = "gpt-3.5-turbo"
MINERU_TOKEN = ""

# 优先级1: config.py
try:
    config = importlib.import_module("config")
    LLM_API_KEY = getattr(config, "LLM_API_KEY", "")
    LLM_API_URL = getattr(config, "LLM_API_URL", LLM_API_URL)
    LLM_MODEL = getattr(config, "LLM_MODEL", LLM_MODEL)
    MINERU_TOKEN = getattr(config, "MINERU_TOKEN", "")
    print("✅ 从 config.py 加载配置")
except ImportError:
    pass

# 优先级2: mykey.py（兼容本地开发环境）
if not LLM_API_KEY:
    try:
        _sys_path_backup = sys.path[:]
        sys.path.insert(0, str(Path(__file__).parent.parent))  # ../mykey.py
        sys.path.insert(0, str(Path(__file__).parent))         # ./mykey.py
        mykey = importlib.import_module("mykey")
        # 优先使用 Mimo_1 配置（本地部署的mimo模型）
        for cfg_name in ["native_oai_config_3", "native_oai_config_2", "native_oai_config_1"]:
            cfg = getattr(mykey, cfg_name, None)
            if cfg and cfg.get("apikey"):
                LLM_API_KEY = cfg["apikey"]
                LLM_API_URL = f"{cfg['apibase']}/chat/completions" if "/chat/completions" not in cfg.get("apibase", "") else cfg["apibase"]
                LLM_MODEL = cfg.get("model", LLM_MODEL)
                print(f"✅ 从 mykey.py.{cfg_name} 加载配置 (模型: {LLM_MODEL})")
                break
        # 加载 MinerU token
        if not MINERU_TOKEN:
            MINERU_TOKEN = getattr(mykey, "MINERU_TOKEN", "")
        sys.path = _sys_path_backup
    except (ImportError, Exception) as e:
        sys.path = _sys_path_backup if '_sys_path_backup' in dir() else sys.path
        pass

if not LLM_API_KEY:
    print("⚠️ 未找到LLM API配置，请创建 config.py 或确保 mykey.py 可用")

MINERU_BASE = "https://mineru.net"

# ─── 欺诈模式知识库加载 ───
FRAUD_PATTERNS_PATH = Path(__file__).parent / "fraud_patterns.json"
FRAUD_PATTERNS = []
PATTERN_HINTS = ""
if FRAUD_PATTERNS_PATH.exists():
    try:
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            pattern_data = json.load(f)
            FRAUD_PATTERNS = pattern_data.get("patterns", [])
        print(f"✅ 加载欺诈模式知识库成功，共{len(FRAUD_PATTERNS)}条检测模式")
        # 构建提示词片段
        PATTERN_HINTS = "\n## 最新欺诈模式知识库（社区贡献+PubPeer案例汇总）\n"
        for idx, p in enumerate(FRAUD_PATTERNS, 1):
            PATTERN_HINTS += f"{idx}. [{p['risk_level']}风险] {p['name']}：{p['detection_hint']}\n"
    except Exception as e:
        print(f"⚠️ 知识库加载失败: {e}，使用默认检测规则")

# ══════════════════════════════════════════════════════════════
# 审查体系配置 - LLM System Prompt
# ══════════════════════════════════════════════════════════════
SYSTEM_PROMPT_TPL = """你是一个严厉的学术论文审查专家（耿同学标准）。
你需要结合以下维度对输入的论文文本进行审查，输出严格的JSON格式：

## 审查维度
1. 数据与结果自洽性 — 数字前后矛盾、统计量不一致、图表数据不匹配
2. 图片与图表异常 — 描述性分析图片可疑特征（旋转复用、背景一致、拼接痕迹）
3. 方法论严谨性 — 样本量不足、缺乏多重比较校正、实验设计缺陷
4. 结构与引用规范性 — 自引率异常、引用质量差、逻辑谬误
5. 作者与期刊可信度 — 产出异常、利益冲突未披露、同行评审缺失

## 检查项（耿同学六式 + 7类红旗）
- 耿同学六式：图片复用/数据造假/图片拼接/统计异常/产出异常/方法矛盾
- 7类红旗：引用质量差/逻辑谬误/方法论缺陷/可疑结论/同行评审缺失/利益冲突未披露/语言质量差

请按以下JSON格式输出（确保JSON合法，无多余内容）：
{{
  "summary": "一句话总评",
  "risk_level": "高/中/低/可疑黑产",
  "detection_score": 0,
  "checks": [
    {{
      "category": "数据与结果/图片与图表/方法论/结构与引用/作者与期刊",
      "item": "检查项名称",
      "verdict": "🚩红旗/⚠️疑点/✅通过",
      "source_text": "必须填写：论文原文中的直接摘录；若无直接证据写'未找到直接原文证据'",
      "evidence": "必须填写：具体证据，引用原文片段并说明所在章节/表图/段落线索",
      "reason": "必须填写：为什么该证据支持此判定，说明可疑逻辑链",
      "detail": "详细分析说明：包含影响范围、需人工复核的点、若为通过也说明依据"
    }}
  ],
  "conclusion": "综合结论与行动建议"
}}
{pattern_hints}
"""

# 动态构建系统提示词
SYSTEM_PROMPT = SYSTEM_PROMPT_TPL.format(pattern_hints=PATTERN_HINTS)

# ══════════════════════════════════════════════════════════════
# MinerU API 模块 — PDF转Markdown
# ══════════════════════════════════════════════════════════════

def _http_request(url, method="GET", headers=None, data=None, timeout=60):
    """通用HTTP请求封装（使用requests，绕过Cloudflare UA检测）"""
    _BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
    if headers is None:
        headers = {}
    headers.setdefault("User-Agent", _BROWSER_UA)
    if method.upper() == "GET":
        resp = requests.get(url, headers=headers, timeout=timeout)
    elif method.upper() == "POST":
        resp = requests.post(url, headers=headers, data=data, timeout=timeout)
    else:
        resp = requests.request(method, url, headers=headers, data=data, timeout=timeout)
    resp.raise_for_status()
    return resp.content, resp.status_code


def mineru_precision_extract_by_url(pdf_url, model_version="vlm", language="ch",
                                     poll_interval=10, poll_timeout=600):
    """🎯 Precision API — 通过URL解析PDF（需要Token，≤200MB/200页）

    流程：POST创建任务 → GET轮询结果 → 下载zip中的Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    print(f"  🎯 [MinerU Precision] 提交URL任务: {pdf_url[:80]}...")

    # 1. 创建提取任务
    create_url = f"{MINERU_BASE}/api/v4/extract/task"
    payload = json.dumps({"url": pdf_url, "model_version": model_version}).encode()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }
    try:
        resp_data, status = _http_request(create_url, "POST", headers, payload, timeout=30)
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"创建任务失败: {e}"}

    if result.get("code") != 0 and not result.get("data", {}).get("batch_id"):
        return None, {"error": f"创建任务返回异常: {result}"}

    batch_id = result.get("data", {}).get("batch_id")
    if not batch_id:
        return None, {"error": f"未获取到batch_id: {result}"}
    print(f"  ✅ 任务已创建: batch_id={batch_id}")

    # 2. 轮询任务状态
    poll_url = f"{MINERU_BASE}/api/v4/extract/task/{batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        task_list = result.get("data", {}).get("task_list", [])
        if not task_list:
            # 单文件模式
            state = result.get("data", {}).get("state", "unknown")
        else:
            state = task_list[0].get("state", "unknown")

        elapsed = int(time.time() - start)

        if state == "done":
            # 获取zip下载链接
            zip_url = task_list[0].get("zip_url") if task_list else result.get("data", {}).get("zip_url")
            if not zip_url:
                return None, {"error": "任务完成但未获取到下载链接"}

            print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
            markdown = _download_zip_and_extract_md(zip_url, output_dir=output_dir, source_name="url_input", batch_id=batch_id)
            if markdown:
                meta = {"source": "mineru_precision", "batch_id": batch_id,
                        "zip_url": zip_url, "zip_saved_dir": str(output_dir) if output_dir else str(_RUN_OUTPUT_DIR) if _RUN_OUTPUT_DIR else None,
                        "model": model_version, "chars": len(markdown)}
                return markdown, meta
            else:
                return None, {"error": "下载或解压zip失败"}

        elif state == "failed":
            err = task_list[0].get("err_msg", "未知") if task_list else "未知"
            return None, {"error": f"任务失败: {err}"}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), batch_id={batch_id}"}


def mineru_extract_file(file_path, model_version="vlm", language="ch",
                        poll_interval=10, poll_timeout=600, output_dir=None):
    """🎯 MinerU v4 本地文件解析（需要Token，≤200MB/200页）

    流程：POST /api/v4/file-urls/batch → PUT上传至OSS → 轮询任务结果 → 下载zip提取Markdown
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    file_size_mb = file_path.stat().st_size / 1024 / 1024
    print(f"  🎯 [MinerU v4] 上传文件: {file_path.name} ({file_size_mb:.1f}MB)")

    if not MINERU_TOKEN:
        return None, {"error": "MINERU_TOKEN未配置，无法使用MinerU API"}

    auth_headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {MINERU_TOKEN}"
    }

    # 1. 获取上传URL
    try:
        batch_payload = json.dumps({
            "enable_formula": True,
            "language": language,
            "layout_model": "doclayout_yolo",
            "enable_table": True,
            "files": [{"name": file_path.name, "is_ocr": True}]
        }).encode()
        resp_data, _ = _http_request(
            f"{MINERU_BASE}/api/v4/file-urls/batch", "POST",
            auth_headers, batch_payload, timeout=30
        )
        result = json.loads(resp_data.decode())
    except Exception as e:
        return None, {"error": f"获取上传URL失败: {e}"}

    if result.get("code") != 0:
        return None, {"error": f"file-urls/batch返回异常: {result}"}

    batch_id = result["data"]["batch_id"]
    file_urls = result["data"]["file_urls"]
    if not file_urls:
        return None, {"error": "未获取到上传URL"}

    upload_url = file_urls[0]
    print(f"  ✅ 获取上传URL成功: batch_id={batch_id}")

    # 2. 上传文件至OSS（注意：预签名URL不能带Authorization header，也不带Content-Type以免签名不匹配）
    try:
        with open(file_path, "rb") as f:
            file_data = f.read()
        resp = requests.put(upload_url, data=file_data, timeout=120)
        resp.raise_for_status()
        print(f"  ✅ 文件已上传至OSS ({file_size_mb:.1f}MB)")
    except Exception as e:
        return None, {"error": f"上传文件到OSS失败: {e}"}

    # 3. 轮询批量解析结果（file-urls/batch 上传后，用 batch_id 查询 extract-results/batch）
    # 经验验证：/api/v4/extract/task/{batch_id} 会返回 task not found，batch_id 应查批量结果接口。
    poll_url = f"{MINERU_BASE}/api/v4/extract-results/batch/{batch_id}"
    start = time.time()
    state_labels = {"processing": "处理中", "queued": "排队中", "pending": "等待中"}

    while time.time() - start < poll_timeout:
        try:
            resp_data, _ = _http_request(poll_url, "GET", auth_headers, timeout=30)
            result = json.loads(resp_data.decode())
        except Exception as e:
            print(f"  ⚠️ 轮询异常: {e}")
            time.sleep(poll_interval)
            continue

        elapsed = int(time.time() - start)
        if result.get("code") != 0:
            print(f"  ⚠️ [{elapsed}s] 查询结果异常: code={result.get('code')} msg={result.get('msg')}")
            time.sleep(poll_interval)
            continue

        data = result.get("data", {}) or {}
        extract_results = data.get("extract_result") or data.get("task_list") or []
        task = extract_results[0] if extract_results else data
        state = task.get("state", "unknown")

        if state == "done":
            zip_url = task.get("full_zip_url") or data.get("full_zip_url")
            if not zip_url:
                return None, {"error": "任务完成但未获取到下载链接", "batch_id": batch_id,
                              "result": result}

            print(f"  ✅ [{elapsed}s] 解析完成，下载Markdown...")
            markdown = _download_zip_and_extract_md(zip_url, output_dir=output_dir, source_name=file_path.name, batch_id=batch_id)
            if markdown:
                return markdown, {"source": "mineru_v4", "batch_id": batch_id,
                                  "zip_url": zip_url, "model": model_version,
                                  "zip_saved_dir": str(output_dir) if output_dir else str(_RUN_OUTPUT_DIR) if _RUN_OUTPUT_DIR else None,
                                  "chars": len(markdown)}
            return None, {"error": "下载或解压zip失败", "batch_id": batch_id,
                          "zip_url": zip_url}

        elif state == "failed":
            err = task.get("err_msg") or data.get("err_msg") or "未知错误"
            return None, {"error": f"任务失败: {err}", "batch_id": batch_id,
                          "result": result}

        label = state_labels.get(state, state)
        print(f"  ⏳ [{elapsed}s] {label}...")
        time.sleep(poll_interval)

    return None, {"error": f"轮询超时({poll_timeout}s), batch_id={batch_id}",
                  "poll_url": poll_url}


def _download_zip_and_extract_md(zip_url, output_dir=None, source_name=None, batch_id=None):
    """下载zip、按输入同目录保存，并提取Markdown文件（纯标准库实现）"""
    try:
        zip_data, _ = _http_request(zip_url, "GET", timeout=60)
        if output_dir or _RUN_OUTPUT_DIR:
            save_mineru_artifacts(zip_url, zip_data, source_name or "mineru", output_dir=output_dir, batch_id=batch_id)
    except Exception as e:
        print(f"  ❌ 下载zip失败: {e}")
        return None

    # 用 zipfile 从内存解析
    import zipfile, io
    try:
        with zipfile.ZipFile(io.BytesIO(zip_data)) as zf:
            # 找到 .md 文件
            md_files = [n for n in zf.namelist() if n.endswith(".md")]
            if not md_files:
                # 降级：找 .txt 或其他文本
                text_files = [n for n in zf.namelist() if n.endswith((".txt", ".mdown", ".markdown"))]
                md_files = text_files
            if not md_files:
                print(f"  ⚠️ zip中未找到Markdown文件: {zf.namelist()[:10]}")
                # 尝试任何非图片文件
                for n in zf.namelist():
                    if not any(n.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".svg", ".html"]):
                        try:
                            content = zf.read(n).decode("utf-8", errors="ignore")
                            if len(content) > 100:
                                return content
                        except:
                            continue
                return None

            # 读取最大的 .md 文件
            best = None
            best_len = 0
            for md_file in md_files:
                content = zf.read(md_file).decode("utf-8", errors="ignore")
                if len(content) > best_len:
                    best = content
                    best_len = len(content)
            return best
    except zipfile.BadZipFile:
        # 不是zip？尝试直接作为文本
        try:
            return zip_data.decode("utf-8", errors="ignore")
        except:
            return None


def mineru_extract(file_path, language="ch", output_dir=None):
    """MinerU统一入口：使用v4 API
    
    本地文件：上传至OSS → 创建任务 → 轮询结果
    URL：直接提交v4/extract/task
    返回：(markdown_text, meta_dict) 或 (None, error_dict)
    """
    file_path = Path(file_path)
    if file_path.exists():
        return mineru_extract_file(file_path, language=language, output_dir=output_dir)
    else:
        # 当作URL处理
        return mineru_precision_extract_by_url(str(file_path), language=language, output_dir=output_dir)


# ══════════════════════════════════════════════════════════════
# 本地统计检测模块
# ══════════════════════════════════════════════════════════════

def benford_analysis(numbers):
    """Benford定律分析：识别异常数字分布（伪造数据首位数字偏离Benford分布）"""
    if len(numbers) < 100:
        return None, "样本不足(需≥100)"
    digits = [str(abs(int(n)))[0] for n in numbers if abs(int(n)) >= 1]
    if not digits:
        return None, "无有效数字"
    counts = collections.Counter(digits)
    total = len(digits)
    expected = {str(d): math.log10(1 + 1/d) * total for d in range(1, 10)}
    deviations = {}
    for d in range(1, 10):
        d_str = str(d)
        actual = counts.get(d_str, 0)
        exp = expected[d_str]
        deviation = abs(actual - exp) / exp
        deviations[d_str] = deviation
    avg_deviation = sum(deviations.values()) / 9
    return avg_deviation, "高偏差⚠️" if avg_deviation > 0.3 else "正常✅"


def extract_all_numbers(text):
    """提取文本中所有数字（排除年份、页码等噪声）"""
    exclude = {2018, 2019, 2020, 2021, 2022, 2023, 2024, 2025, 2026, 2027,
               1, 2, 3, 4, 5, 6, 7, 8, 9, 10}
    nums = []
    for match in re.finditer(r'\b(\d+\.?\d*)\b', text):
        try:
            n = float(match.group(1))
            if n not in exclude and n > 0:
                nums.append(n)
        except:
            pass
    return nums


def local_stat_check(text):
    """本地统计检测，无需LLM

    包含：Benford定律分析、p值异常检测、标准差异常、数字自洽性
    """
    result = {
        "benford_deviation": None,
        "benford_status": None,
        "p_value_count": 0,
        "p_value_abnormal": 0,
        "p_value_details": [],
        "sd_count": 0,
        "sd_abnormal": 0,
        "number_count": 0,
        "number_consistency": None,
    }
    # 提取数字
    nums = extract_all_numbers(text)
    result["number_count"] = len(nums)
    # Benford分析
    if nums:
        dev, status = benford_analysis(nums)
        result["benford_deviation"] = dev
        result["benford_status"] = status
    # p值检测
    p_matches = re.findall(r'p\s*[=<]\s*(\d+\.?\d*)', text, re.IGNORECASE)
    result["p_value_count"] = len(p_matches)
    for p in p_matches:
        try:
            pv = float(p)
            if pv > 0.05:
                result["p_value_abnormal"] += 1
                result["p_value_details"].append(f"p={'<='+p if pv<=0.001 else '='+p}")
        except:
            pass
    # 标准差异常
    sd_matches = re.findall(r'(?:std|sd|标准差|SE|SEM)\s*[=:≈]\s*(\d+\.?\d*)', text, re.IGNORECASE)
    result["sd_count"] = len(sd_matches)
    # 数字自洽性检查：提取"n=XX"样本量，检查是否有矛盾
    n_matches = re.findall(r'(?:n|N|sample|样本)\s*[=:]\s*(\d+)', text, re.IGNORECASE)
    if len(set(n_matches)) > 1:
        result["number_consistency"] = f"检测到不同样本量: {set(n_matches)}"
    return result


# ══════════════════════════════════════════════════════════════
# 目录级综合分析模块
# ══════════════════════════════════════════════════════════════

def find_project_files(root_path: Path) -> Tuple[Dict, List[Path]]:
    """递归扫描目录，识别论文项目相关的所有文件
    返回：(文件分类字典, 所有有效文件列表)
    """
    SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".md"}
    SUPPLEMENT_KEYWORDS = {"supplement", "supp", "补充材料", "原始数据", "data", "source", "appendix"}
    
    file_categories = {
        "main_paper": None,
        "supplements": [],
        "data_files": [],
        "other": []
    }
    all_files = []
    
    for root, _, files in os.walk(root_path):
        for file in files:
            fpath = Path(root) / file
            ext = fpath.suffix.lower()
            if ext not in SUPPORTED_EXTS:
                continue
            
            fname = fpath.name.lower()
            all_files.append(fpath)
            
            # 分类
            if ext == ".pdf" and (file_categories["main_paper"] is None or 
                                 ("main" in fname or "paper" in fname or "article" in fname)):
                file_categories["main_paper"] = fpath
            elif any(kw in fname for kw in SUPPLEMENT_KEYWORDS):
                file_categories["supplements"].append(fpath)
            elif ext in {".xlsx", ".xlsm", ".csv"}:
                file_categories["data_files"].append(fpath)
            else:
                file_categories["other"].append(fpath)
    
    # 未找到明确主论文则取最大的PDF
    if file_categories["main_paper"] is None:
        pdf_files = [f for f in all_files if f.suffix.lower() == ".pdf"]
        if pdf_files:
            file_categories["main_paper"] = sorted(pdf_files, key=lambda x: x.stat().st_size, reverse=True)[0]
    
    return file_categories, all_files


def extract_text_from_file(file_path: Path, max_chars_per_file=None, use_mineru=False, mineru_lang="ch", output_dir=None) -> str:
    """从任意支持的文件类型中提取文本

    max_chars_per_file=None 表示不截断。目录级分析默认应先完整提取每个文件，
    再由后续 smart_chunk_text 做全目录分块/合并审查，避免“目录里只分析到一个文件/每文件只取开头”。
    """
    ext = file_path.suffix.lower()
    header = f"=== 文件: {file_path.name} ==="
    text = f"\n\n{header}"
    limit = max_chars_per_file if max_chars_per_file is not None else 999999999
    
    try:
        if ext == ".pdf":
            if use_mineru:
                print(f"  ⚙️  使用MinerU API提取PDF全文内容...")
                md, meta = mineru_extract(file_path, language=mineru_lang, output_dir=output_dir)
                if md:
                    text += "\n" + md
                else:
                    err = meta.get("error", "未知错误") if isinstance(meta, dict) else "未知错误"
                    print(f"  ⚠️  MinerU提取失败: {err}，降级为本地PDF提取")
                    pdf_text, _, _ = extract_pdf_text(file_path, max_chars=limit)
                    text += "\n" + pdf_text
            else:
                pdf_text, _, _ = extract_pdf_text(file_path, max_chars=limit)
                text += "\n" + pdf_text
        elif ext == ".docx" and DOCX_SUPPORTED:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += "\n" + para.text
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join([cell.text for cell in row.cells])
        elif ext in {".xlsx", ".xlsm"} and EXCEL_SUPPORTED:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                text += f"\n[工作表: {sheet_name}]"
                sheet = wb[sheet_name]
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if max_chars_per_file is not None and i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    row_str = " | ".join([str(v) for v in row if v is not None])
                    if row_str.strip():
                        text += "\n" + row_str
        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if max_chars_per_file is not None and i > 1000:
                        text += "\n[数据过多，已截断]"
                        break
                    text += "\n" + " | ".join(row)
        elif ext in {".txt", ".md"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text += "\n" + f.read(limit)
    except Exception as e:
        text += f"\n[文件解析失败: {str(e)}]"
    
    if max_chars_per_file is not None and len(text) > max_chars_per_file + len(header) + 4:
        return text[:max_chars_per_file + len(header) + 4] + "\n[文本过长已截断]"
    return text


# ══════════════════════════════════════════════════════════════
# PDF原始提取模块（MinerU不可用时的降级方案）
# ══════════════════════════════════════════════════════════════

def extract_pdf_text(filepath, max_chars=8000):
    """从PDF文件中提取文本（纯标准库实现，MinerU的降级方案）"""
    with open(filepath, "rb") as f:
        raw = f.read()
    parts = []
    for s in re.findall(rb"stream\r?\n(.*?)\r?\nendstream", raw, re.DOTALL):
        try:
            dec = zlib.decompress(s)
            for x in re.findall(rb"\((.*?)\)\s*Tj", dec):
                d = x.decode("latin-1", errors="ignore")
                if len(d.strip()) > 1:
                    parts.append(d)
            for bt in re.findall(rb"BT(.*?)ET", dec, re.DOTALL):
                for x in re.findall(rb"\((.*?)\)", bt):
                    d = x.decode("latin-1", errors="ignore")
                    if len(d.strip()) > 1:
                        parts.append(d)
        except:
            pass
    text = re.sub(r"\s+", " ", " ".join(parts)).strip()
    meta = {"size_mb": round(len(raw) / 1024 / 1024, 2), "total_chars": len(text),
            "extraction_method": "raw_pdf_stream"}
    return text[:max_chars], meta, raw


# ══════════════════════════════════════════════════════════════
# LLM调用模块
# ══════════════════════════════════════════════════════════════

def smart_chunk_text(text, chunk_size=8000, overlap=1000):
    """智能分块：按段落边界切割，保留重叠区确保上下文连贯。

    关键保证：任何返回块都不会超过 chunk_size。MinerU Markdown 中常出现超长表格/图片
    JSON/参考文献段落，旧实现会把单个超长段落直接塞进一个块，导致 LLM 网关 504。
    返回: [(chunk_text, chunk_index, total_chunks), ...]
    """
    if chunk_size <= 0:
        chunk_size = 8000
    overlap = max(0, min(overlap, chunk_size // 4))

    def hard_split(s, size):
        """把单段硬切到 size 内，优先在换行/句号附近切。"""
        parts = []
        s = s or ""
        while len(s) > size:
            cut = size
            window_start = max(0, size - 800)
            candidates = [s.rfind(sep, window_start, size) for sep in ["\n", "。", ". ", "; ", "；", ", ", "，", " "]]
            best = max(candidates)
            if best > max(200, size // 2):
                cut = best + 1
            parts.append(s[:cut].strip())
            s = s[cut:].strip()
        if s.strip():
            parts.append(s.strip())
        return parts

    if len(text) <= chunk_size:
        return [(text, 0, 1)]

    # 按双换行分段，并先硬切超长段落，确保后续组装不会产生超限块
    raw_paragraphs = re.split(r'\n{2,}', text)
    paragraphs = []
    for p in raw_paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(p) > chunk_size:
            paragraphs.extend(hard_split(p, chunk_size))
        else:
            paragraphs.append(p)

    chunks = []
    current = ""
    for para in paragraphs:
        # 防御：如果仍有超长段，继续硬切
        para_parts = hard_split(para, chunk_size) if len(para) > chunk_size else [para]
        for part in para_parts:
            candidate = (current + "\n\n" + part).strip() if current else part
            if current and len(candidate) > chunk_size:
                chunks.append(current[:chunk_size])
                prefix = current[-overlap:] + "\n\n" if overlap > 0 and len(current) > overlap else ""
                candidate = (prefix + part).strip()
                if len(candidate) > chunk_size:
                    # overlap + part 仍超限时，放弃 overlap，硬切 part
                    for sub in hard_split(part, chunk_size):
                        if len(sub) <= chunk_size:
                            chunks.append(sub)
                    current = ""
                else:
                    current = candidate
            else:
                current = candidate
    if current:
        chunks.append(current[:chunk_size])

    # 最终保险：展开任何异常超限块
    safe_chunks = []
    for c in chunks:
        if len(c) > chunk_size:
            safe_chunks.extend(hard_split(c, chunk_size))
        elif c.strip():
            safe_chunks.append(c)

    total = len(safe_chunks)
    return [(c, i, total) for i, c in enumerate(safe_chunks)]


def call_llm(text, max_retries=None, chunk_info=None, timeout=None):
    """调用OpenAI兼容API进行语义审查。

    改进点：
    - timeout/max_retries可由CLI控制，便于不稳定网关降级；
    - 每次失败打印尝试编号，日志更容易判断是否卡死；
    - 去除重复payload构建，减少维护歧义。
    """
    if max_retries is None:
        max_retries = int(globals().get("LLM_RETRIES", 1))
    if timeout is None:
        timeout = int(globals().get("LLM_TIMEOUT", 45))

    if chunk_info and chunk_info[1] > 1:
        idx, total = chunk_info
        user_msg = (
            f"审查以下论文文本（第{idx+1}/{total}段，请重点关注本段内容，"
            f"同时注意与其他段落的逻辑连贯性）：\n\n{text}"
        )
    else:
        user_msg = f"审查以下论文文本：\n\n{text}"

    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg}
        ],
        "temperature": 0.2
    }
    _BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
    last_err = None
    for attempt in range(max_retries + 1):
        try:
            if attempt:
                print(f"     ↻ API重试 {attempt}/{max_retries}（timeout={timeout}s）")
            resp = requests.post(
                LLM_API_URL, json=payload,
                headers={"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json", "User-Agent": _BROWSER_UA},
                timeout=timeout
            )
            resp.raise_for_status()
            result = resp.json()
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            last_err = e
            if attempt < max_retries:
                print(f"     ⚠️ API尝试 {attempt+1}/{max_retries+1} 失败: {str(e)[:160]}")
                time.sleep(3 * (attempt + 1))
            else:
                raise RuntimeError(f"API调用失败({attempt+1}次, timeout={timeout}s): {str(last_err)[:180]}...")


def merge_chunk_reports(reports, stat_result=None):
    """合并多块审查结果：去重、合并检查项、重新评估风险等级

    reports: [parse_report返回的dict, ...]
    """
    if len(reports) == 1:
        return reports[0]

    # 1. 收集所有检查项，按 (category, item) 去重
    seen_keys = set()
    all_checks = []
    for i, r in enumerate(reports):
        if r.get("parse_error"):
            continue
        for c in r.get("checks", []):
            key = (c.get("category", ""), c.get("item", ""))
            if key not in seen_keys:
                seen_keys.add(key)
                c["_source_chunk"] = i + 1
                all_checks.append(c)
            else:
                # 合并：如果已有同key项，补充来源信息
                for existing in all_checks:
                    ekey = (existing.get("category", ""), existing.get("item", ""))
                    if ekey == key:
                        # 保留更严重的判定
                        severity = {"🚩红旗": 3, "⚠️疑点": 2, "✅通过": 1}
                        old_s = severity.get(existing.get("verdict", ""), 0)
                        new_s = severity.get(c.get("verdict", ""), 0)
                        if new_s > old_s:
                            existing["verdict"] = c["verdict"]
                            existing["evidence"] = c.get("evidence", existing.get("evidence", ""))
                            existing["detail"] = c.get("detail", existing.get("detail", ""))
                        # 补充证据
                        if c.get("evidence") and c["evidence"] not in existing.get("evidence", ""):
                            existing["evidence"] = (existing.get("evidence", "") + 
                                                     f" [第{c.get('_source_chunk', i+1)}段补充: {c['evidence']}]")
                        break

    # 2. 统计红旗/疑点数量
    red_flags = sum(1 for c in all_checks if "红旗" in c.get("verdict", ""))
    warnings = sum(1 for c in all_checks if "疑点" in c.get("verdict", ""))

    # 3. 重新评估风险等级
    if red_flags >= 3:
        risk_level = "高"
    elif red_flags >= 1 or warnings >= 3:
        risk_level = "中"
    elif warnings >= 1:
        risk_level = "低"
    else:
        risk_level = "低"

    # 结合统计结果调整
    if stat_result:
        if stat_result.get("benford_deviation") and stat_result["benford_deviation"] > 0.3:
            risk_level = "高" if risk_level == "中" else risk_level
        if stat_result.get("p_value_abnormal", 0) > 2:
            risk_level = max(risk_level, "中", key=["低", "中", "高", "可疑黑产"].index)

    # 4. 计算打假得分
    detection_score = red_flags * 30 + warnings * 10
    if stat_result and stat_result.get("benford_deviation"):
        detection_score += int(stat_result["benford_deviation"] * 50)

    # 5. 综合所有summary
    summaries = [r.get("summary", "") for r in reports if not r.get("parse_error")]
    merged_summary = " | ".join(s for s in summaries if s)[:200]
    if len(summaries) > 1:
        merged_summary = f"[合并{len(reports)}段审查] {merged_summary}"

    # 6. 综合conclusion
    conclusions = [r.get("conclusion", "") for r in reports if not r.get("parse_error") and r.get("conclusion")]
    merged_conclusion = "\n\n".join(conclusions) if conclusions else ""

    # 清理临时字段
    for c in all_checks:
        c.pop("_source_chunk", None)

    return {
        "summary": merged_summary,
        "risk_level": risk_level,
        "detection_score": detection_score,
        "checks": all_checks,
        "conclusion": merged_conclusion,
        "_merged_from": len(reports),
    }


# ══════════════════════════════════════════════════════════════
# 报告解析与格式化
# ══════════════════════════════════════════════════════════════

def parse_report(content):
    """解析LLM返回的JSON报告，容错处理"""
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass
    m = re.search(r'\{[\s\S]*\}', content)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            pass
    return {"raw_output": content, "parse_error": True}


def _md_escape_cell(text):
    """Markdown表格单元格转义与压缩。"""
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    return text.replace("|", "\\|")


def _brief_text(text, limit=180):
    """压缩长文本，保留报告可读性。"""
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) > limit:
        return text[:limit].rstrip() + "…"
    return text


def _is_suspicious_check(c):
    verdict = str(c.get("verdict", ""))
    return ("红旗" in verdict) or ("疑点" in verdict) or ("可疑" in verdict)


def _check_source_text(c):
    """尽量提取LLM给出的原文支撑/证据字段。兼容不同JSON字段名。"""
    for k in ("source_text", "quote", "original_text", "原文", "原文摘录", "evidence"):
        v = c.get(k)
        if isinstance(v, (list, tuple)):
            v = "；".join(str(x) for x in v if x)
        if v:
            return str(v)
    return ""


def _check_reason(c):
    """提取可疑原因/细节，优先detail，其次reason/explanation。"""
    for k in ("detail", "reason", "analysis", "explanation", "说明"):
        v = c.get(k)
        if v:
            return str(v)
    return ""


def format_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为Markdown报告"""
    risk_icons = {"高": "🔴", "中": "🟡", "低": "🟢", "可疑黑产": "⚫️"}
    lines = [
        f"# 📄 学术论文审查报告 [耿同学标准]",
        f"",
        f"**文件**: `{pdf_path}`",
        f"**文件大小**: {meta.get('size_mb', 'N/A')} MB",
        f"**提取字符数**: {meta.get('total_chars', meta.get('chars', 'N/A'))}",
        f"**提取方式**: {meta.get('extraction_method', meta.get('source', 'N/A'))}",
    ]
    # 显示分块信息（如果是分块审查）
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        lines.append(f"**审查方式**: 分块审查 | {meta['chunk_count']}块 | 单块上限{meta['chunk_size']}字符 | 重叠{meta['overlap']}字符")
    if meta.get("llm_coverage"):
        failed_chunks = meta.get("llm_failed_chunks") or []
        if meta.get("llm_partial_report") or failed_chunks:
            lines.append(f"**LLM覆盖率**: ⚠️ 部分报告，仅成功审查 {meta.get('llm_coverage')} 个分块；失败块: {failed_chunks or '无'}")
            lines.append("> ⚠️ 本报告只基于成功返回的LLM分块合并，未覆盖失败分块全文；结论只能作为阶段性结果，建议稍后用 `--llm-cache-only` 或更稳定API补跑。")
        else:
            lines.append(f"**LLM覆盖率**: ✅ {meta.get('llm_coverage')} 个分块全部成功")
    lines.append(f"**审查时间**: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    lines.extend([
        f"",
        f"## 📊 本地统计检测结果",
        f"| 检测项 | 结果 | 状态 |",
        f"|--------|------|------|",
        f"| Benford分布偏差 | {round(stat_result['benford_deviation'],3) if stat_result['benford_deviation'] else '样本不足'} | {stat_result['benford_status'] or 'N/A'} |",
        f"| p值数量/异常 | {stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个>0.05 | {'⚠️异常' if stat_result['p_value_abnormal'] else '✅正常'} |",
        f"| 标准差提及 | {stat_result['sd_count']}处 | N/A |",
    ])
    lines.append(f"| 提取数字数 | {stat_result['number_count']} | - |")

    if stat_result.get("number_consistency"):
        lines.append(f"| 数字自洽性 | {stat_result['number_consistency']} | ⚠️矛盾 |")

    lines.append("")

    if report.get("parse_error"):
        lines.append("## ⚠️ LLM报告解析失败（原始输出）")
        lines.append(f"```\n{report['raw_output']}\n```")
        return "\n".join(lines)

    lines.append(f"## 总评: {report.get('summary', 'N/A')}")
    risk = report.get('risk_level', '未知')
    lines.append(f"**风险等级**: {risk_icons.get(risk, '⚪')} {risk}")
    lines.append(f"**打假得分**: {report.get('detection_score', 0)} (越高越可疑)")
    lines.append("")

    checks = report.get("checks", [])
    if checks:
        suspicious = [c for c in checks if _is_suspicious_check(c)]
        lines.append("## 🚩 可疑点证据汇总表")
        lines.append("")
        if suspicious:
            lines.append("| # | 风险判定 | 分类/检查项 | 原文证据摘录 | 可疑原因 |")
            lines.append("|---|----------|-------------|--------------|----------|")
            for i, c in enumerate(suspicious, 1):
                cat_item = f"{c.get('category', 'N/A')} / {c.get('item', 'N/A')}"
                lines.append(
                    f"| {i} | {_md_escape_cell(c.get('verdict', 'N/A'))} | {_md_escape_cell(cat_item)} | "
                    f"{_md_escape_cell(_brief_text(_check_source_text(c), 220) or '未提供明确原文摘录')} | "
                    f"{_md_escape_cell(_brief_text(_check_reason(c), 220) or '未提供详细原因')} |"
                )
        else:
            lines.append("> 未发现红旗/疑点项；仍建议人工核验关键数据、图表和引用。")
        lines.append("")

        lines.append("## 🔍 全部检查项概览")
        lines.append("")
        lines.append("| # | 分类 | 检查项 | 判定 | 证据摘要 |")
        lines.append("|---|------|--------|------|----------|")
        for i, c in enumerate(checks, 1):
            lines.append(f"| {i} | {_md_escape_cell(c.get('category', 'N/A'))} | {_md_escape_cell(c.get('item', 'N/A'))} | {_md_escape_cell(c.get('verdict', 'N/A'))} | {_md_escape_cell(_brief_text(_check_source_text(c), 120) or '-')} |")
        lines.append("")

        lines.append("## 📋 逐条详细分析（含原文支撑）")
        lines.append("")
        for i, c in enumerate(checks, 1):
            lines.append(f"### {i}. {c.get('category', 'N/A')} - {c.get('item', 'N/A')} — {c.get('verdict', 'N/A')}")
            source = _check_source_text(c)
            reason = _check_reason(c)
            if source:
                lines.append(f"> **原文/证据摘录**: {source}")
            else:
                lines.append("> **原文/证据摘录**: LLM未提供明确原文摘录，请人工回查对应段落。")
            if reason:
                lines.append(f"\n**可疑原因/详细说明**：{reason}")
            lines.append("")

    if report.get("conclusion"):
        lines.append("## 📝 综合结论")
        lines.append(f"\n{report['conclusion']}")

    return "\n".join(lines)


def format_html_report(report, pdf_path, meta, stat_result):
    """将审查结果格式化为美观的HTML报告"""
    risk_colors = {"高": "#dc2626", "中": "#f59e0b", "低": "#16a34a", "可疑黑产": "#7c3aed"}
    risk_icons = {"高": "🔴", "中": "🟡", "低": "🟢", "可疑黑产": "⚫️"}
    risk = report.get('risk_level', '未知')
    risk_color = risk_colors.get(risk, "#6b7280")
    risk_icon = risk_icons.get(risk, "⚪")

    # 统计检测状态
    benford_val = round(stat_result['benford_deviation'], 3) if stat_result['benford_deviation'] else '样本不足'
    benford_status = stat_result.get('benford_status', 'N/A') or 'N/A'
    p_abnormal = stat_result['p_value_abnormal']
    p_status_class = "status-warn" if p_abnormal else "status-ok"

    # 分块信息
    chunk_info = ""
    if meta.get("chunk_count") and meta["chunk_count"] > 1:
        chunk_info = f"""
        <div class="info-row">
            <span class="info-label">审查方式</span>
            <span class="info-value">分块审查 | {meta['chunk_count']}块 | 单块上限{meta['chunk_size']}字符 | 重叠{meta['overlap']}字符</span>
        </div>"""

    # 数字自洽性
    number_consistency = ""
    if stat_result.get("number_consistency"):
        number_consistency = f"""
        <tr>
            <td>数字自洽性</td>
            <td>{stat_result['number_consistency']}</td>
            <td><span class="status-warn">⚠️ 矛盾</span></td>
        </tr>"""

    # LLM覆盖率/部分报告提示
    coverage_banner = ""
    if meta.get("llm_coverage"):
        failed_chunks = meta.get("llm_failed_chunks") or []
        if meta.get("llm_partial_report") or failed_chunks:
            coverage_banner = f"""
  <div class="section coverage-warning">
    <h2>⚠️ 部分报告：LLM覆盖不足</h2>
    <p><strong>成功审查分块</strong>: {_html_escape(meta.get('llm_coverage'))}</p>
    <p><strong>失败块</strong>: {_html_escape(failed_chunks or '无')}</p>
    <p>本报告只基于成功返回的LLM分块合并，未覆盖失败分块全文；结论只能作为阶段性结果。建议稍后使用 <code>--llm-cache-only</code> 复用成功缓存，或切换更稳定API补跑。</p>
  </div>"""
        else:
            coverage_banner = f"""
  <div class="section coverage-ok">
    <h2>✅ LLM覆盖率</h2>
    <p>{_html_escape(meta.get('llm_coverage'))} 个分块全部成功。</p>
  </div>"""

    # 解析失败
    if report.get("parse_error"):
        checks_html = f"""
        <div class="section">
            <h2>⚠️ LLM报告解析失败（原始输出）</h2>
            <pre class="error-block">{_html_escape(report.get('raw_output', ''))}</pre>
        </div>"""
        conclusion_html = ""
    else:
        # 可疑点证据汇总表 + 全部检查概览 + 详细分析
        checks = report.get("checks", [])
        suspicious = [c for c in checks if _is_suspicious_check(c)]

        suspicious_rows = ""
        for i, c in enumerate(suspicious, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = "verdict-red" if "红旗" in verdict else ("verdict-yellow" if ("疑点" in verdict or "可疑" in verdict) else "verdict-green")
            cat_item = f"{c.get('category', 'N/A')} / {c.get('item', 'N/A')}"
            suspicious_rows += f"""
            <tr>
                <td>{i}</td>
                <td><span class="{verdict_class}">{_html_escape(verdict)}</span></td>
                <td>{_html_escape(cat_item)}</td>
                <td class="evidence-cell">{_html_escape(_brief_text(_check_source_text(c), 260) or '未提供明确原文摘录')}</td>
                <td class="reason-cell">{_html_escape(_brief_text(_check_reason(c), 260) or '未提供详细原因')}</td>
            </tr>"""
        if not suspicious_rows:
            suspicious_rows = """
            <tr><td colspan="5" class="muted">未发现红旗/疑点项；仍建议人工核验关键数据、图表和引用。</td></tr>"""

        checks_table_rows = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = "verdict-red" if "红旗" in verdict else ("verdict-yellow" if ("疑点" in verdict or "可疑" in verdict) else "verdict-green")
            checks_table_rows += f"""
            <tr>
                <td>{i}</td>
                <td>{_html_escape(c.get('category', 'N/A'))}</td>
                <td>{_html_escape(c.get('item', 'N/A'))}</td>
                <td><span class="{verdict_class}">{_html_escape(verdict)}</span></td>
                <td class="evidence-cell">{_html_escape(_brief_text(_check_source_text(c), 140) or '-')}</td>
            </tr>"""

        detail_cards = ""
        for i, c in enumerate(checks, 1):
            verdict = c.get('verdict', 'N/A')
            verdict_class = "verdict-red" if "红旗" in verdict else ("verdict-yellow" if ("疑点" in verdict or "可疑" in verdict) else "verdict-green")
            source = _check_source_text(c)
            reason = _check_reason(c)
            detail_cards += f"""
            <div class="detail-card">
                <div class="detail-header">
                    <span class="detail-num">#{i}</span>
                    <span class="detail-cat">{_html_escape(c.get('category', 'N/A'))}</span>
                    <span class="detail-item">{_html_escape(c.get('item', 'N/A'))}</span>
                    <span class="{verdict_class} detail-verdict">{_html_escape(verdict)}</span>
                </div>
                <div class="detail-evidence"><strong>原文/证据摘录</strong><blockquote>{_html_escape(source or 'LLM未提供明确原文摘录，请人工回查对应段落。')}</blockquote></div>
                <div class="detail-text"><strong>可疑原因/详细说明</strong><p>{_html_escape(reason or 'LLM未提供详细说明。')}</p></div>
            </div>"""

        checks_html = f"""
        <div class="section evidence-summary">
            <h2>🚩 可疑点证据汇总表</h2>
            <p class="section-hint">优先阅读本表：每条可疑点均尽量给出原文摘录和判断理由，便于快速人工复核。</p>
            <table class="checks-table evidence-table">
                <thead><tr><th>#</th><th>风险判定</th><th>分类/检查项</th><th>原文证据摘录</th><th>可疑原因</th></tr></thead>
                <tbody>{suspicious_rows}</tbody>
            </table>
        </div>
        <div class="section">
            <h2>🔍 全部检查项概览</h2>
            <table class="checks-table">
                <thead><tr><th>#</th><th>分类</th><th>检查项</th><th>判定</th><th>证据摘要</th></tr></thead>
                <tbody>{checks_table_rows}</tbody>
            </table>
        </div>
        <div class="section">
            <h2>📋 逐条详细分析（含原文支撑）</h2>
            {detail_cards}
        </div>"""

        conclusion_html = ""
        if report.get("conclusion"):
            conclusion_html = f"""
            <div class="section conclusion-section">
                <h2>📝 综合结论</h2>
                <p class="conclusion-text">{_html_escape(report['conclusion'])}</p>
            </div>"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>📄 学术论文审查报告</title>
<style>
  :root {{
    --bg: #0f172a;
    --surface: #1e293b;
    --surface2: #334155;
    --text: #e2e8f0;
    --text-muted: #94a3b8;
    --accent: #38bdf8;
    --border: #475569;
    --red: #ef4444;
    --yellow: #f59e0b;
    --green: #22c55e;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Noto Sans SC', sans-serif;
    background: var(--bg);
    color: var(--text);
    line-height: 1.6;
    padding: 20px;
  }}
  .container {{ max-width: 960px; margin: 0 auto; }}
  .header {{
    background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 32px;
    margin-bottom: 24px;
    text-align: center;
  }}
  .header h1 {{ font-size: 28px; margin-bottom: 16px; }}
  .risk-badge {{
    display: inline-block;
    font-size: 24px;
    font-weight: 700;
    padding: 8px 24px;
    border-radius: 999px;
    color: #fff;
    background: {risk_color};
    margin: 8px 0;
  }}
  .meta-grid {{
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 8px 24px;
    text-align: left;
    margin-top: 16px;
    font-size: 14px;
    color: var(--text-muted);
  }}
  .meta-grid strong {{ color: var(--text); }}
  .section {{
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 24px;
    margin-bottom: 20px;
  }}
  .section h2 {{
    font-size: 20px;
    margin-bottom: 16px;
    padding-bottom: 8px;
    border-bottom: 1px solid var(--border);
  }}
  table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
  th, td {{ padding: 10px 12px; text-align: left; border-bottom: 1px solid var(--border); }}
  th {{ background: var(--surface2); color: var(--text-muted); font-weight: 600; text-transform: uppercase; font-size: 12px; }}
  tr:hover {{ background: rgba(56,189,248,0.05); }}
  .status-ok {{ color: var(--green); font-weight: 600; }}
  .status-warn {{ color: var(--yellow); font-weight: 600; }}
  .verdict-red {{ color: var(--red); font-weight: 700; }}
  .verdict-yellow {{ color: var(--yellow); font-weight: 700; }}
  .verdict-green {{ color: var(--green); font-weight: 700; }}
  .detail-card {{
    background: var(--surface2);
    border-radius: 8px;
    padding: 16px;
    margin-bottom: 12px;
    border-left: 4px solid var(--accent);
  }}
  .detail-header {{ display: flex; align-items: center; gap: 12px; flex-wrap: wrap; }}
  .section-hint {{ color: var(--text-muted); margin: -6px 0 14px; font-size: 14px; }}
  .evidence-summary {{ border-left: 4px solid var(--red); }}
  .checks-table {{ table-layout: fixed; }}
  .checks-table th:nth-child(1), .checks-table td:nth-child(1) {{ width: 44px; text-align: center; }}
  .checks-table th:nth-child(2), .checks-table td:nth-child(2) {{ width: 120px; }}
  .evidence-table th:nth-child(4), .evidence-table td:nth-child(4),
  .evidence-table th:nth-child(5), .evidence-table td:nth-child(5) {{ width: 28%; }}
  .evidence-cell, .reason-cell {{
    color: var(--text);
    line-height: 1.55;
    word-break: break-word;
    overflow-wrap: anywhere;
  }}
  .muted {{ color: var(--text-muted); text-align: center; padding: 18px; }}
  .detail-num {{
    background: var(--accent);
    color: var(--bg);
    border-radius: 50%;
    width: 28px; height: 28px;
    display: flex; align-items: center; justify-content: center;
    font-weight: 700; font-size: 13px;
  }}
  .detail-cat {{ color: var(--accent); font-weight: 600; }}
  .detail-item {{ flex: 1; }}
  .detail-verdict {{ font-size: 14px; }}
  .detail-evidence {{
    margin-top: 10px;
    padding: 10px 14px;
    background: rgba(245,158,11,0.1);
    border-radius: 6px;
    font-size: 14px;
    color: #fbbf24;
  }}
  .detail-text {{
    margin-top: 8px;
    font-size: 14px;
    color: var(--text-muted);
    white-space: pre-wrap;
  }}
  .conclusion-section {{ border-left: 4px solid var(--green); }}
  .conclusion-text {{ font-size: 16px; white-space: pre-wrap; color: var(--text); }}
  .coverage-warning {{
    border-left: 5px solid var(--yellow);
    background: rgba(245,158,11,0.10);
  }}
  .coverage-warning code {{ background: var(--surface2); padding: 2px 6px; border-radius: 4px; }}
  .coverage-ok {{
    border-left: 5px solid var(--green);
    background: rgba(16,185,129,0.08);
  }}
  .error-block {{
    background: rgba(239,68,68,0.1);
    border: 1px solid var(--red);
    border-radius: 8px;
    padding: 16px;
    white-space: pre-wrap;
    font-family: monospace;
    font-size: 13px;
    color: #fca5a5;
    overflow-x: auto;
  }}
  .score-bar {{
    height: 8px;
    background: var(--surface2);
    border-radius: 4px;
    overflow: hidden;
    margin-top: 8px;
  }}
  .score-fill {{
    height: 100%;
    border-radius: 4px;
    background: linear-gradient(90deg, var(--green), var(--yellow), var(--red));
    transition: width 0.5s;
  }}
  .footer {{
    text-align: center;
    color: var(--text-muted);
    font-size: 12px;
    margin-top: 32px;
    padding: 16px;
  }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>📄 学术论文审查报告</h1>
    <div class="risk-badge">{risk_icon} {risk}</div>
    <div style="margin-top:8px; color:var(--text-muted);">
      打假得分: <strong style="color:{risk_color}">{report.get('detection_score', 0)}</strong>（越高越可疑）
      <div class="score-bar"><div class="score-fill" style="width:{min(report.get('detection_score', 0), 100)}%"></div></div>
    </div>
    <div class="meta-grid">
      <div><strong>文件</strong>: {_html_escape(pdf_path)}</div>
      <div><strong>文件大小</strong>: {meta.get('size_mb', 'N/A')} MB</div>
      <div><strong>提取字符数</strong>: {meta.get('total_chars', meta.get('chars', 'N/A'))}</div>
      <div><strong>提取方式</strong>: {meta.get('extraction_method', meta.get('source', 'N/A'))}</div>
      {chunk_info.replace('class="info-row"', '').replace('class="info-label"', '').replace('class="info-value"', '').strip().replace('<span>', '<div>').replace('</span>', '</div>') if chunk_info else '<div></div><div></div>'}
      <div><strong>审查时间</strong>: {time.strftime('%Y-%m-%d %H:%M:%S')}</div>
    </div>
  </div>

  {coverage_banner}

  <div class="section">
    <h2>📊 本地统计检测结果</h2>
    <table>
      <thead><tr><th>检测项</th><th>结果</th><th>状态</th></tr></thead>
      <tbody>
        <tr><td>Benford分布偏差</td><td>{benford_val}</td><td>{benford_status}</td></tr>
        <tr><td>p值数量/异常</td><td>{stat_result['p_value_count']} / {stat_result['p_value_abnormal']}个&gt;0.05</td><td><span class="{p_status_class}">{'⚠️异常' if p_abnormal else '✅正常'}</span></td></tr>
        <tr><td>标准差提及</td><td>{stat_result['sd_count']}处</td><td>N/A</td></tr>
        <tr><td>提取数字数</td><td>{stat_result['number_count']}</td><td>-</td></tr>
        {number_consistency}
      </tbody>
    </table>
  </div>

  {checks_html}
  {conclusion_html}

  <div class="footer">
    Generated by <strong>Veritas</strong> — 学术论文自动审查工具（耿同学标准） | {time.strftime('%Y-%m-%d %H:%M:%S')}
  </div>
</div>
</body>
</html>"""
    return html


def _html_escape(text):
    """HTML特殊字符转义"""
    if not text:
        return ""
    return (str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\n", "<br>"))

def update_patterns(comments_file):
    """从PubPeer评论文本中用LLM提取新的欺诈模式，更新知识库
    
    comments_file: 包含PubPeer评论文本的文件路径
    """
    from datetime import datetime
    
    comments_path = Path(comments_file)
    if not comments_path.exists():
        print(f"❌ 评论文本文件不存在: {comments_path}")
        return 1
    
    with open(comments_path, "r", encoding="utf-8") as f:
        comments_text = f.read()
    
    if len(comments_text.strip()) < 20:
        print("❌ 评论文本内容过少，请提供更完整的PubPeer评论内容")
        return 1
    
    print(f"📖 已读取评论文本: {len(comments_text)}字符")
    print("🤖 正在用LLM分析评论，提取欺诈模式...")
    
    # 构建提取prompt
    extract_prompt = f"""分析以下来自PubPeer的学术评论，提取其中涉及的学术论文造假/可疑手法。

要求：
1. 每个造假手法提取为一个独立的模式条目
2. 按JSON数组格式输出，每个条目包含：id(英文大写下划线), category(分类), name(中文名), description(详细描述), detection_hint(检测提示), risk_level(高/中/低)
3. 只提取确实存在的造假手法，不要臆造
4. 合并相似的造假手法

PubPeer评论内容：
{comments_text}

输出格式：
[
  {{
    "id": "PATTERN_ID",
    "category": "图片与图表/数据与结果/方法论/结构与引用/作者与期刊",
    "name": "手法名称",
    "description": "手法描述",
    "detection_hint": "审查时如何检测此手法",
    "risk_level": "高/中/低"
  }}
]"""
    
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": "你是一个学术论文打假专家，擅长从PubPeer评论中识别和归纳造假手法。"},
            {"role": "user", "content": extract_prompt}
        ],
        "temperature": 0.3,
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LLM_API_KEY}",
    }
    
    req = urllib.request.Request(
        LLM_API_URL,
        data=json.dumps(payload).encode(),
        headers=headers,
        method="POST",
    )
    
    try:
        resp = urllib.request.urlopen(req, timeout=60)
        result = json.loads(resp.read().decode("utf-8"))
        content = result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"❌ LLM调用失败: {e}")
        return 1
    
    # 解析LLM输出的JSON
    json_match = re.search(r'\[[\s\S]*\]', content)
    if not json_match:
        print("❌ LLM未能输出有效的JSON格式，请重试")
        print(f"原始输出: {content[:500]}")
        return 1
    
    try:
        new_patterns = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        print(f"❌ JSON解析失败: {e}")
        return 1
    
    if not new_patterns:
        print("⚠️ 未能从评论中提取到新的欺诈模式")
        return 0
    
    # 加载现有知识库
    if FRAUD_PATTERNS_PATH.exists():
        with open(FRAUD_PATTERNS_PATH, "r", encoding="utf-8") as f:
            kb_data = json.load(f)
        existing_ids = {p["id"] for p in kb_data.get("patterns", [])}
    else:
        kb_data = {"schema_version": "1.0", "last_updated": "", "contributors": ["community"], "patterns": []}
        existing_ids = set()
    
    # 去重合并
    added = 0
    for p in new_patterns:
        if p.get("id") and p["id"] not in existing_ids:
            kb_data["patterns"].append(p)
            existing_ids.add(p["id"])
            added += 1
            print(f"  ✅ 新增: [{p.get('risk_level','?')}] {p.get('name','?')}")
        else:
            print(f"  ⏭️ 跳过已存在: {p.get('name','?')} ({p.get('id','?')})")
    
    if added > 0:
        kb_data["last_updated"] = datetime.now().strftime("%Y-%m-%d")
        with open(FRAUD_PATTERNS_PATH, "w", encoding="utf-8") as f:
            json.dump(kb_data, f, ensure_ascii=False, indent=2)
        print(f"\n🎉 知识库已更新！新增{added}条模式，总计{len(kb_data['patterns'])}条")
    else:
        print("\n⚠️ 无新增模式，知识库未变更")
    
    return 0


# ══════════════════════════════════════════════════════════════
# 腾讯朱雀AI文本检测辅助功能
# ══════════════════════════════════════════════════════════════

ZHUQUE_URL = "https://matrix.tencent.com/ai-detect/"


def copy_to_clipboard(text: str) -> bool:
    """跨平台复制文本到系统剪贴板"""
    system = platform.system()
    try:
        if system == "Windows":
            # Windows: 使用clip命令
            process = subprocess.Popen(
                ["clip"], stdin=subprocess.PIPE, shell=True
            )
            process.communicate(text.encode("utf-16"))
            return process.returncode == 0
        elif system == "Darwin":  # macOS
            process = subprocess.Popen(
                ["pbcopy"], stdin=subprocess.PIPE
            )
            process.communicate(text.encode("utf-8"))
            return process.returncode == 0
        else:  # Linux
            # 优先尝试xclip，其次xsel
            for cmd in [["xclip", "-selection", "clipboard"], ["xsel", "--clipboard", "--input"]]:
                try:
                    process = subprocess.Popen(cmd, stdin=subprocess.PIPE)
                    process.communicate(text.encode("utf-8"))
                    if process.returncode == 0:
                        return True
                except FileNotFoundError:
                    continue
            return False
    except Exception as e:
        print(f"⚠️ 剪贴板写入失败: {e}")
        return False


def launch_zhuque_ai_detect(text: str):
    """启动腾讯朱雀AI文本检测：复制文本到剪贴板 → 打开检测页面 → 弹窗提醒"""
    print("\n" + "=" * 60)
    print("🤖 腾讯朱雀AI文本检测")
    print("=" * 60)

    # 1) 复制文本到剪贴板
    # 朱雀检测有字数限制，截取前8000字符
    detect_text = text[:8000]
    if len(text) > 8000:
        print(f"⚠️ 文本较长({len(text)}字符)，仅复制前8000字符到剪贴板（朱雀字数限制）")

    clip_ok = copy_to_clipboard(detect_text)
    if clip_ok:
        print("✅ 文本已复制到剪贴板")
    else:
        print("❌ 剪贴板写入失败，请手动复制论文文本")

    # 2) 打开浏览器
    print(f"🌐 正在打开朱雀AI检测页面...")
    webbrowser.open(ZHUQUE_URL)

    # 3) 弹窗提醒
    system = platform.system()
    try:
        if system == "Windows":
            import ctypes
            ctypes.windll.user32.MessageBoxW(
                0,
                "论文文本已复制到剪贴板！\n\n"
                "请在打开的朱雀AI检测页面中粘贴文本并点击检测。\n"
                "检测完成后，点击确定继续后续审查流程。",
                "🤖 朱雀AI文本检测",
                0x40  # MB_ICONINFORMATION
            )
        elif system == "Darwin":
            subprocess.run([
                "osascript", "-e",
                'display dialog "论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。" '
                'buttons {"确定"} default button "确定" with title "🤖 朱雀AI文本检测" with icon note'
            ])
        else:  # Linux
            # 尝试zenity
            try:
                subprocess.run([
                    "zenity", "--info", "--title=🤖 朱雀AI文本检测", "--width=400",
                    "--text=论文文本已复制到剪贴板！\n\n请在朱雀AI检测页面中粘贴文本并点击检测。\n检测完成后，点击确定继续后续审查流程。"
                ])
            except FileNotFoundError:
                # 降级为终端提示
                input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")
    except Exception:
        # 最终降级：终端等待
        input("\n⏸️ 论文文本已复制到剪贴板，请在浏览器中粘贴检测。\n检测完成后按回车继续...")

    print("✅ 朱雀AI检测流程结束，继续后续审查...")


# ──────────────────────────────────────────────────────────────
# AI图片检测（imagedetector.com）
# ──────────────────────────────────────────────────────────────

IMAGE_DETECT_URL = "https://imagedetector.com/"

# 支持的图片扩展名
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".tif", ".webp"}


def extract_images_from_pdf(pdf_path: str) -> List[str]:
    """从PDF中提取内嵌图片到临时目录，返回图片路径列表

    优先使用PyMuPDF(fitz)，降级使用pdf2image整页渲染
    """
    images = []
    tmp_dir = os.path.join(os.path.dirname(pdf_path), "_veritas_images_tmp")
    os.makedirs(tmp_dir, exist_ok=True)

    # 方案1：PyMuPDF提取内嵌图片
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        img_count = 0
        for page_idx in range(len(doc)):
            page = doc[page_idx]
            img_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(img_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        ext = base_image.get("ext", "png")
                        if ext not in ("png", "jpg", "jpeg", "bmp", "tiff", "webp"):
                            ext = "png"
                        fname = f"page{page_idx + 1}_img{img_idx + 1}.{ext}"
                        fpath = os.path.join(tmp_dir, fname)
                        with open(fpath, "wb") as f:
                            f.write(base_image["image"])
                        # 过滤掉太小的图片（图标、装饰等）
                        if os.path.getsize(fpath) > 5000:
                            images.append(fpath)
                            img_count += 1
                except Exception:
                    continue
        doc.close()
        if img_count > 0:
            print(f"  📎 PyMuPDF提取 {img_count} 张内嵌图片 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ PyMuPDF提取失败: {e}")

    # 方案2：pdf2image整页渲染
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(pdf_path, dpi=200)
        for i, page_img in enumerate(pages):
            fname = f"page{i + 1}_full.png"
            fpath = os.path.join(tmp_dir, fname)
            page_img.save(fpath, "PNG")
            if os.path.getsize(fpath) > 10000:
                images.append(fpath)
        if images:
            print(f"  📎 pdf2image渲染 {len(images)} 页 → {tmp_dir}")
            return images
    except ImportError:
        pass
    except Exception as e:
        print(f"  ⚠️ pdf2image渲染失败: {e}")

    return images


def collect_image_files(input_path: str) -> List[str]:
    """收集论文相关图片文件（从目录扫描 + PDF提取）

    返回所有图片的绝对路径列表
    """
    images = []
    p = Path(input_path)

    if p.is_file() and p.suffix.lower() == ".pdf":
        # PDF文件：提取内嵌图片
        print("  📸 从PDF中提取图片...")
        extracted = extract_images_from_pdf(str(p))
        images.extend(extracted)
    elif p.is_dir():
        # 目录：扫描所有图片文件
        for ext in IMAGE_EXTENSIONS:
            for f in p.rglob(f"*{ext}"):
                if f.stat().st_size > 5000:  # 过滤小图标
                    images.append(str(f))

    return images


def launch_image_ai_detect(input_path: str):
    """启动imagedetector.com图片AI检测：收集图片 → 打开浏览器 → 逐张提醒上传"""
    print("\n" + "=" * 60)
    print("🖼️ AI图片检测 (imagedetector.com)")
    print("=" * 60)

    # 1) 收集图片
    images = collect_image_files(input_path)

    if not images:
        print("⚠️ 未找到可检测的图片文件")
        print("💡 提示: 可手动访问 https://imagedetector.com/ 上传图片检测")
        return

    # 限制数量（太多图片会非常耗时）
    MAX_IMAGES = 30
    if len(images) > MAX_IMAGES:
        print(f"⚠️ 图片数量过多({len(images)}张)，仅处理前{MAX_IMAGES}张（小图标已过滤）")
        images = images[:MAX_IMAGES]

    print(f"  📋 共 {len(images)} 张图片待检测")

    # 2) 打开浏览器
    print(f"🌐 正在打开 imagedetector.com 检测页面...")
    webbrowser.open(IMAGE_DETECT_URL)
    time.sleep(2)  # 等待浏览器加载

    # 3) 逐张提醒上传
    system = platform.system()
    for idx, img_path in enumerate(images, 1):
        img_name = os.path.basename(img_path)
        msg = (
            f"图片 [{idx}/{len(images)}]: {img_name}\n\n"
            f"路径: {img_path}\n\n"
            f"请在 imagedetector.com 页面中上传此图片并点击检测。\n"
            f"检测完成后，点击确定继续下一张图片。"
        )

        print(f"\n  📸 [{idx}/{len(images)}] {img_name}")
        print(f"     路径: {img_path}")

        try:
            if system == "Windows":
                import ctypes
                result = ctypes.windll.user32.MessageBoxW(
                    0,
                    msg,
                    f"🖼️ AI图片检测 ({idx}/{len(images)})",
                    0x41  # MB_OKCANCEL + MB_ICONWARNING
                )
                if result == 2:  # IDCANCEL
                    print("  ⏹️ 用户取消图片检测流程")
                    break
            elif system == "Darwin":
                subprocess.run([
                    "osascript", "-e",
                    f'display dialog "{msg}" '
                    f'buttons {{"取消", "确定"}} default button "确定" with title "🖼️ AI图片检测 ({idx}/{len(images)})" with icon note'
                ])
            else:  # Linux
                try:
                    subprocess.run([
                        "zenity", "--info", "--title=🖼️ AI图片检测", "--width=450",
                        f"--text={msg}"
                    ])
                except FileNotFoundError:
                    input(f"  ⏸️ 按回车继续下一张...")
        except Exception:
            input(f"  ⏸️ 检测完成后按回车继续下一张...")

    print("✅ AI图片检测流程结束，继续后续审查...")

    # 清理临时提取目录
    tmp_dir = os.path.join(os.path.dirname(input_path) if os.path.isfile(input_path) else input_path, "_veritas_images_tmp")
    if os.path.isdir(tmp_dir):
        try:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
            print("  🧹 已清理临时图片目录")
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser(
        description="学术论文自动审查工具（耿同学标准 + MinerU）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # MinerU提取 + 完整审查（推荐）
  python paper_audit.py paper.pdf --mineru

  # 仅原始PDF文本提取 + 审查
  python paper_audit.py paper.pdf

  # 指定输出路径
  python paper_audit.py paper.pdf --mineru -o report.md --json
  
  # 更新欺诈模式知识库（从PubPeer评论自动提取新pattern）
  python paper_audit.py --update-patterns pubpeer_comments.txt
""")
    parser.add_argument("pdf_path", nargs='?', help="待审查的文件路径或论文目录路径（支持PDF/Word/Excel/Supplement等，更新模式下无需提供）")
    parser.add_argument("--update-patterns", metavar="COMMENTS_FILE", 
                        help="从PubPeer评论文本文件中自动提取新的欺诈模式，更新知识库")
    parser.add_argument("--mineru", action="store_true",
                        help="使用MinerU API将PDF转为Markdown再审查（推荐，质量更高）")
    parser.add_argument("--mineru-model", default="vlm",
                        choices=["pipeline", "vlm", "MinerU-HTML"],
                        help="MinerU模型版本（默认vlm，仅Precision API生效）")
    parser.add_argument("--mineru-lang", default="ch",
                        help="MinerU OCR语言（默认ch=中英，en=英文，japan=日文）")
    parser.add_argument("--no-mineru", action="store_true",
                        help="强制使用原始PDF文本提取，禁用MinerU")
    parser.add_argument("--max-chars", type=int, default=4096,
                        help="LLM分块单块最大字符数（默认4096；超过4096会自动压到4096）")
    parser.add_argument("--output", "-o", help="输出报告文件路径（默认输出到同目录）")
    parser.add_argument("--json", action="store_true", help="同时保存原始JSON结果")
    parser.add_argument("--ai-detect", action="store_true", help="开启腾讯朱雀AI文本检测：自动复制文本到剪贴板+打开检测页面")
    parser.add_argument("--image-detect", action="store_true", help="开启AI图片检测（imagedetector.com）：自动提取图片+打开检测页面+逐张提醒上传")
    parser.add_argument("--llm-provider", default=None,
                        help="从mykey.py指定LLM供应商配置名，例如 native_oai_config_3；会覆盖config.py里的LLM，但不影响MinerU配置")
    parser.add_argument("--ignore-config-llm", action="store_true",
                        help="忽略config.py中的LLM设置，强制使用mykey.py的LLM配置（推荐用于避开旧fufu接口）")
    parser.add_argument("--no-resume", action="store_true",
                        help="禁用断点续作缓存，强制重新提取文本和重新LLM审查")
    parser.add_argument("--llm-timeout", type=int, default=45,
                        help="单次LLM请求超时时间秒数（默认45；不稳定网关可调小以更快跳过）")
    parser.add_argument("--llm-retries", type=int, default=1,
                        help="每次LLM调用内部重试次数（默认1，即最多2次尝试）")
    parser.add_argument("--strict-failed-chunks", action="store_true",
                        help="严格模式：失败块首轮+补跑仍失败时停止生成报告；默认生成覆盖率受限的部分报告")
    parser.add_argument("--llm-cache-only", action="store_true",
                        help="只复用已有成功LLM分块缓存生成部分报告，不再调用API；适合接口不稳定时先查看已有结果")
    args = parser.parse_args()
    if getattr(args, "max_chars", 4096) > 4096:
        print(f"⚠️ --max-chars={args.max_chars} 超过4096，已自动调整为4096")
        args.max_chars = 4096
    if getattr(args, "max_chars", 4096) < 512:
        print(f"⚠️ --max-chars={args.max_chars} 过小，已自动调整为512")
        args.max_chars = 512
    global LLM_TIMEOUT, LLM_RETRIES
    LLM_TIMEOUT = max(10, int(getattr(args, "llm_timeout", 45) or 45))
    LLM_RETRIES = max(0, int(getattr(args, "llm_retries", 1) or 0))

    # ─── 知识库更新模式 ───
    if args.update_patterns:
        return update_patterns(args.update_patterns)

    # ─── 正常审查模式 ───
    if not args.pdf_path:
        parser.error("审查模式需要提供path参数（文件或目录，或使用 --update-patterns 更新知识库）")

    input_path = Path(args.pdf_path)
    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        return 1

    output_dir, output_stem = get_output_base(input_path)
    setup_run_logging(input_path)
    print(f"📁 所有输出将保存到: {output_dir}")
    if args.llm_provider or args.ignore_config_llm:
        load_mykey_llm_config(args.llm_provider)
    print(f"🤖 当前LLM: model={LLM_MODEL}, url={LLM_API_URL}")
    resume_dir = get_resume_dir(output_dir, output_stem)
    if args.no_resume:
        print("♻️ 已禁用断点续作缓存，本次将重新执行提取和LLM审查")
    else:
        print(f"🔁 断点续作缓存目录: {resume_dir}")
    resume_event(resume_dir, "init", "done", f"input={input_path}; llm={LLM_MODEL}; url={LLM_API_URL}; max_chars={args.max_chars}")
    progress_bar(0, 5, "初始化完成")

    # ─── 阶段1：文本提取（支持单个文件/整个论文目录） ───
    extract_cache_path = resume_dir / "stage1_extract.json"
    cached_extract = None if args.no_resume else _json_load(extract_cache_path)
    if cached_extract and cached_extract.get("input") == str(input_path.resolve()) and cached_extract.get("use_mineru") == (args.mineru and not args.no_mineru):
        full_text = cached_extract.get("full_text", "")
        meta = cached_extract.get("meta", {})
        raw_pdf = None
        use_mineru = cached_extract.get("use_mineru", args.mineru and not args.no_mineru)
        print(f"🔁 断点续作：复用阶段1文本缓存 {extract_cache_path} ({len(full_text)}字符)")
        resume_event(resume_dir, "stage1_extract", "cache_hit", f"chars={len(full_text)}", cache=str(extract_cache_path))
        progress_bar(1, 5, "阶段1/5 文本提取缓存命中")
    else:
        full_text = None
    meta = {}
    raw_pdf = None
    use_mineru = args.mineru and not args.no_mineru

    if full_text is None and input_path.is_dir():
        print(f"📂 检测到输入为目录，正在扫描所有论文相关文件...")
        file_classes, all_files = find_project_files(input_path)
        print(f"✅ 找到 {len(all_files)} 个相关文件:")
        for cat, files in file_classes.items():
            if not files:
                continue
            if isinstance(files, Path):
                print(f"  - {cat}: {files.name}")
            else:
                print(f"  - {cat}: {len(files)} 个文件")
        
        # 提取所有文件文本合并
        full_text = ""
        total_files = len(all_files)
        for idx, file_path in enumerate(all_files, 1):
            print(f"  📝 提取 [{idx}/{total_files}] {file_path.name}...")
            progress_bar(idx - 1, total_files, f"阶段1/5 提取文件: {file_path.name}")
            file_content = extract_text_from_file(file_path, max_chars_per_file=None,
                                                  use_mineru=use_mineru,
                                                  mineru_lang=args.mineru_lang,
                                                  output_dir=output_dir)
            progress_bar(idx, total_files, f"阶段1/5 已完成: {file_path.name}")
            full_text += f"\n\n=== 文件: {file_path.name} 路径: {file_path.relative_to(input_path)} ==="
            full_text += "\n" + file_content
        
        def _class_count(v):
            if v is None:
                return 0
            if isinstance(v, Path):
                return str(v)
            try:
                return len(v)
            except TypeError:
                return str(v)

        meta = {
            "input_type": "directory",
            "total_files": total_files,
            "file_classes": {k: _class_count(v) for k, v in file_classes.items()},
            "total_chars": len(full_text),
            "extractor": "directory_multi_format"
        }
        print(f"\n✅ 所有文件提取完成，总长度: {len(full_text)} 字符")
        progress_bar(1, 5, "阶段1/5 文本提取完成")
    elif full_text is None:
        # 单个文件走原有流程
        pdf_path = input_path
        print(f"📄 检测到输入为单个文件: {pdf_path.name}")

        if use_mineru:
            print(f"📡 [MinerU] 正在将PDF转为Markdown: {pdf_path.name}")
            md_text, md_meta = mineru_extract(pdf_path, language=args.mineru_lang, output_dir=output_dir)
            if md_text:
                full_text = md_text  # 保留全文
                meta = {
                    "size_mb": round(pdf_path.stat().st_size / 1024 / 1024, 2),
                    "total_chars": len(md_text),
                    "chars_sent": len(md_text),
                    "extraction_method": f"mineru_{md_meta.get('source', 'unknown')}",
                }
                if md_meta.get("batch_id"):
                    meta["mineru_batch_id"] = md_meta["batch_id"]
                if md_meta.get("task_id"):
                    meta["mineru_task_id"] = md_meta["task_id"]
                print(f"✅ MinerU提取完成: {len(md_text)} 字符（全文保留）")
                progress_bar(1, 5, "阶段1/5 MinerU文本提取完成")
            else:
                err = md_meta.get("error", "未知错误") if md_meta else "未知错误"
                print(f"❌ MinerU提取失败: {err}")
                print(f"⚠️ 降级使用原始PDF文本提取...")
                use_mineru = False

        if not use_mineru or full_text is None:
            print(f"📖 正在提取PDF文本: {pdf_path}")
            # extract_pdf_text的max_chars参数传大值以获取全文
            full_text, meta, raw_pdf = extract_pdf_text(str(pdf_path), max_chars=999999)
            if not full_text:
                print("❌ 未能从PDF中提取到文本（可能是扫描件或加密PDF）")
                print("💡 建议: 使用 --mineru 参数通过MinerU API提取（支持OCR）")
                return 1
            print(f"✅ 提取完成: {meta['total_chars']} 字符（全文保留）")
            progress_bar(1, 5, "阶段1/5 PDF文本提取完成")

    if not args.no_resume and full_text:
        _json_save(extract_cache_path, {
            "input": str(input_path.resolve()),
            "use_mineru": use_mineru,
            "mineru_lang": args.mineru_lang,
            "full_text": full_text,
            "meta": meta,
            "saved_at": time.strftime("%F %T"),
        })
        resume_event(resume_dir, "stage1_extract", "saved", f"chars={len(full_text)}; use_mineru={use_mineru}", cache=str(extract_cache_path))

    # ─── 朱雀AI文本检测（可选） ───
    if args.ai_detect:
        launch_zhuque_ai_detect(full_text)

    # ─── AI图片检测（可选） ───
    if args.image_detect:
        launch_image_ai_detect(str(input_path))

    # ─── 阶段2：本地统计检测（使用全文，统计不截断） ───
    progress_bar(1, 5, "阶段2/5 开始本地统计检测")
    print(f"🔢 正在执行本地统计检测...")
    stat_result = local_stat_check(full_text)
    benford_str = f"{round(stat_result['benford_deviation'],3)}" if stat_result['benford_deviation'] else 'N/A'
    print(f"✅ 统计检测完成: Benford偏差={benford_str}, p值异常={stat_result['p_value_abnormal']}, 数字数={stat_result['number_count']}")
    resume_event(resume_dir, "stage2_stat", "done", f"numbers={stat_result['number_count']}; benford={benford_str}")
    progress_bar(2, 5, "阶段2/5 本地统计检测完成")

    # ─── 阶段3：智能分块 + LLM语义审查（冗余机制） ───
    chunk_size = min(int(args.max_chars), 4096)  # LLM单块硬上限4096字符
    overlap = min(512, chunk_size // 8)  # 重叠区约12.5%，最多512字符

    chunks = smart_chunk_text(full_text, chunk_size=chunk_size, overlap=overlap)
    total_chunks = len(chunks)
    llm_cache_key = _text_fingerprint(full_text, f"{LLM_API_URL}|{LLM_MODEL}|{chunk_size}|{overlap}")
    llm_cache_dir = resume_dir / f"llm_{llm_cache_key}"
    llm_cache_dir.mkdir(parents=True, exist_ok=True)
    resume_event(resume_dir, "stage3_llm", "start", f"chunks={total_chunks}; chunk_size={chunk_size}; overlap={overlap}", cache_dir=str(llm_cache_dir))

    progress_bar(2, 5, f"阶段3/5 开始LLM审查：{total_chunks}块")

    if total_chunks == 1:
        # 短论文：直接全文审查
        print(f"🔍 论文长度({len(full_text)}字符)在单块范围内，直接审查...")
        single_cache = llm_cache_dir / "chunk_0000.json"
        cached = None if args.no_resume else _json_load(single_cache)
        if cached:
            print(f"🔁 断点续作：复用LLM审查缓存 {single_cache}")
            resume_event(resume_dir, "stage3_llm_chunk", "cache_hit", "chunk=1/1", cache=str(single_cache))
            report = cached.get("report", {"parse_error": True, "raw_output": "缓存格式异常"})
        else:
            try:
                raw_content = call_llm(full_text)
                report = parse_report(raw_content)
                if not args.no_resume:
                    _json_save(single_cache, {"report": report, "raw_content": raw_content, "saved_at": time.strftime("%F %T")})
                    resume_event(resume_dir, "stage3_llm_chunk", "saved", "chunk=1/1", cache=str(single_cache))
            except Exception as e:
                print(f"❌ LLM调用失败: {e}")
                report = {"parse_error": True, "raw_output": str(e)}
    else:
        # 长论文：分块审查 + 合并
        print(f"🔍 论文较长({len(full_text)}字符)，分为{total_chunks}块(每块≤{chunk_size}字符，重叠{overlap}字符)进行审查...")
        chunk_reports = [None] * total_chunks
        failed_chunks = []

        def _run_chunk_once(chunk_text, chunk_idx, retry=False):
            chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
            print(("  🔁 重试" if retry else "  📝 审查") + f"第{chunk_idx+1}/{total_chunks}块({len(chunk_text)}字符)...")
            raw_content = call_llm(chunk_text, chunk_info=(chunk_idx, total_chunks))
            chunk_report = parse_report(raw_content)
            if chunk_report.get("parse_error"):
                raise RuntimeError(f"LLM返回解析失败: {str(chunk_report.get('raw_output',''))[:180]}")
            if not args.no_resume:
                _json_save(chunk_cache, {"report": chunk_report, "raw_content": raw_content, "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "ok", "retry": retry})
                resume_event(resume_dir, "stage3_llm_chunk", "retry_saved" if retry else "saved", f"chunk={chunk_idx+1}/{total_chunks}; chars={len(chunk_text)}", cache=str(chunk_cache))
            return chunk_report

        for chunk_text, chunk_idx, _ in chunks:
            progress_bar(chunk_idx, total_chunks, f"阶段3/5 LLM审查中：第{chunk_idx+1}/{total_chunks}块")
            chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
            cached = None if args.no_resume else _json_load(chunk_cache)
            if cached and cached.get("status") == "ok" and cached.get("report") and not cached.get("report", {}).get("parse_error"):
                print(f"     ↳ 断点续作：复用第{chunk_idx+1}块成功LLM缓存")
                resume_event(resume_dir, "stage3_llm_chunk", "cache_hit", f"chunk={chunk_idx+1}/{total_chunks}", cache=str(chunk_cache))
                chunk_reports[chunk_idx] = cached.get("report")
            elif getattr(args, "llm_cache_only", False):
                print(f"     ↳ cache-only：第{chunk_idx+1}块无成功缓存，跳过API调用")
                failed_chunks.append((chunk_text, chunk_idx, "cache_only_no_success_cache"))
                resume_event(resume_dir, "stage3_llm_chunk", "cache_only_miss", f"chunk={chunk_idx+1}/{total_chunks}", cache=str(chunk_cache))
            else:
                try:
                    chunk_reports[chunk_idx] = _run_chunk_once(chunk_text, chunk_idx, retry=False)
                except Exception as e:
                    print(f"  ⚠️ 第{chunk_idx+1}块LLM调用/解析失败，先记录并继续其他块: {e}")
                    failed_chunks.append((chunk_text, chunk_idx, str(e)))
                    if not args.no_resume:
                        _json_save(chunk_cache, {"report": {"parse_error": True, "raw_output": str(e)}, "raw_content": str(e), "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "failed_pending_retry"})
                        resume_event(resume_dir, "stage3_llm_chunk", "failed_pending_retry", f"chunk={chunk_idx+1}/{total_chunks}; error={e}", cache=str(chunk_cache))
            if chunk_reports[chunk_idx] and not chunk_reports[chunk_idx].get("parse_error"):
                print(f"     → 第{chunk_idx+1}块风险: {chunk_reports[chunk_idx].get('risk_level', '未知')}")
            progress_bar(chunk_idx + 1, total_chunks, f"阶段3/5 LLM审查完成：第{chunk_idx+1}/{total_chunks}块")

        if failed_chunks:
            print(f"🔁 首轮完成，按顺序重试失败块: {[idx+1 for _, idx, _ in failed_chunks]}")
            resume_event(resume_dir, "stage3_llm_retry", "start", f"failed_chunks={[idx+1 for _, idx, _ in failed_chunks]}; cache_only={getattr(args, 'llm_cache_only', False)}")
            still_failed = []
            if getattr(args, "llm_cache_only", False):
                still_failed = [(idx, first_error) for _, idx, first_error in failed_chunks]
                print("⚠️ cache-only模式：不调用API重试，直接用已有成功缓存生成部分报告。")
            else:
                for chunk_text, chunk_idx, first_error in failed_chunks:
                    try:
                        chunk_reports[chunk_idx] = _run_chunk_once(chunk_text, chunk_idx, retry=True)
                        print(f"     ✅ 第{chunk_idx+1}块重试成功")
                    except Exception as e:
                        print(f"     ❌ 第{chunk_idx+1}块重试仍失败: {e}")
                        still_failed.append((chunk_idx, str(e)))
                        chunk_cache = llm_cache_dir / f"chunk_{chunk_idx:04d}.json"
                        if not args.no_resume:
                            _json_save(chunk_cache, {"report": {"parse_error": True, "raw_output": str(e)}, "raw_content": str(e), "saved_at": time.strftime("%F %T"), "chunk_index": chunk_idx, "total_chunks": total_chunks, "status": "failed_final", "first_error": first_error})
                            resume_event(resume_dir, "stage3_llm_chunk", "failed_final", f"chunk={chunk_idx+1}/{total_chunks}; error={e}", cache=str(chunk_cache))
            if still_failed:
                failed_nums = [idx + 1 for idx, _ in still_failed]
                detail = "; ".join([f"第{idx+1}块: {err}" for idx, err in still_failed])
                resume_event(resume_dir, "stage3_llm_retry", "still_failed", f"still_failed={failed_nums}; strict={args.strict_failed_chunks}")
                if args.strict_failed_chunks:
                    raise RuntimeError("LLM分块重试后仍失败，严格模式停止生成报告: " + detail)
                print(f"⚠️ 仍有{len(still_failed)}块失败，将生成【部分报告】并在报告中标注覆盖不足: {failed_nums}")
            else:
                resume_event(resume_dir, "stage3_llm_retry", "done", "all failed chunks recovered")

        successful_count = sum(1 for r in chunk_reports if r is not None and not r.get("parse_error"))
        failed_final = []
        for idx in range(total_chunks):
            if chunk_reports[idx] is None or chunk_reports[idx].get("parse_error"):
                failed_final.append(idx + 1)
        meta["llm_success_chunks"] = successful_count
        meta["llm_failed_chunks"] = failed_final
        meta["llm_coverage"] = f"{successful_count}/{total_chunks}"
        meta["llm_partial_report"] = bool(failed_final)

        chunk_reports = [r for r in chunk_reports if r is not None and not r.get("parse_error")]
        if not chunk_reports:
            report = {
                "parse_error": True,
                "raw_output": f"所有LLM分块均失败，无法生成语义审查报告。失败块: {failed_final}。可换稳定API后断点续跑。"
            }
            resume_event(resume_dir, "stage4_merge", "skipped_no_success", report["raw_output"])
            progress_bar(4, 5, "阶段4/5 无成功LLM块，仅生成失败说明报告")
        else:

            progress_bar(3, 5, "阶段3/5 LLM审查完成")
            # 合并所有块的审查结果
            progress_bar(3, 5, "阶段4/5 开始合并审查结果")
            print(f"🔗 正在合并{len(chunk_reports)}块审查结果...")
            report = merge_chunk_reports(chunk_reports, stat_result)
            if report.get("_merged_from"):
                print(f"✅ 合并完成: 来自{report['_merged_from']}块，共{len(report.get('checks', []))}个检查项")
            meta["chunk_count"] = total_chunks
            meta["chunk_size"] = chunk_size
            meta["overlap"] = overlap
            if meta.get("llm_partial_report"):
                warning = f"注意：本报告仅覆盖 {meta.get('llm_coverage')} 个LLM分块；失败块: {meta.get('llm_failed_chunks')}。结论不完整，建议换稳定API后断点续跑。"
                report["_partial_warning"] = warning
                report["summary"] = warning + " " + str(report.get("summary", ""))
            resume_event(resume_dir, "stage4_merge", "done", f"checks={len(report.get('checks', [])) if isinstance(report, dict) else 'N/A'}; coverage={meta.get('llm_coverage')}")
            progress_bar(4, 5, "阶段4/5 审查结果合并完成")

    # ─── 阶段5：生成报告 ───
    progress_bar(4, 5, "阶段5/5 开始生成报告")
    report_input = str(input_path)
    md_report = format_report(report, report_input, meta, stat_result)

    # 生成HTML报告
    html_report = format_html_report(report, report_input, meta, stat_result)

    # 确定输出路径（优先HTML）
    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = output_dir / output_path
        # 如果用户指定了.md后缀，同时生成HTML
        html_output_path = output_path.with_suffix(".html")
    else:
        if input_path.is_dir():
            output_path = input_path / "audit_report.audit.md"
            html_output_path = input_path / "audit_report.audit.html"
        else:
            output_path = input_path.with_suffix(".audit.md")
            html_output_path = input_path.with_suffix(".audit.html")

    # 写入Markdown报告
    output_path.write_text(md_report, encoding="utf-8")
    print(f"✅ Markdown报告已保存: {output_path}")
    resume_event(resume_dir, "stage5_report", "markdown_saved", str(output_path))

    # 写入HTML报告
    html_output_path.write_text(html_report, encoding="utf-8")
    print(f"✅ HTML报告已保存: {html_output_path}")
    resume_event(resume_dir, "stage5_report", "html_saved", str(html_output_path))

    if args.json:
        if input_path.is_dir():
            json_path = input_path / "audit_report.audit.json"
        else:
            json_path = input_path.with_suffix(".audit.json")
        json_path.write_text(
            json.dumps({"llm_report": report, "stat_result": stat_result, "meta": meta},
                       ensure_ascii=False, indent=2),
            encoding="utf-8")
        print(f"✅ 原始JSON已保存: {json_path}")

    resume_event(resume_dir, "all", "done", "audit completed")
    progress_bar(5, 5, "阶段5/5 全部完成")
    print(f"🧾 完整日志: {_RUN_LOG_FILE}")

    # 自动打开HTML报告
    try:
        html_abs = str(html_output_path.resolve())
        webbrowser.open(f"file:///{html_abs}" if platform.system() == "Windows" else f"file://{html_abs}")
        print(f"🌐 已在浏览器中打开HTML报告")
    except Exception as e:
        print(f"⚠️ 自动打开浏览器失败: {e}，请手动打开: {html_output_path}")

    # 打印摘要
    if not report.get("parse_error"):
        risk = report.get("risk_level", "未知")
        print(f"\n📊 风险等级: {risk} | 总评: {report.get('summary', 'N/A')}")

    return 0


if __name__ == "__main__":
    exit(main())
